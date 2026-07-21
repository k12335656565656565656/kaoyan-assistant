import io
import sys
import unittest
from pathlib import Path


PROJECT_DIR = Path(__file__).resolve().parents[1]
if str(PROJECT_DIR) not in sys.path:
    sys.path.insert(0, str(PROJECT_DIR))

from docx import Document

from services.knowledge_memorization_docx_service import build_knowledge_memorization_docx


def _document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


class KnowledgeMemorizationDocxServiceTests(unittest.TestCase):
    def test_kb_deployment_installs_docx_runtime(self):
        requirements = (PROJECT_DIR / "requirements_kb.txt").read_text(encoding="utf-8")
        self.assertIn("python-docx", requirements)
        self.assertIn("lxml", requirements)
        self.assertIn("python-dotenv", requirements)

    def test_builds_clean_a4_student_memorization_docx(self):
        points = [
            {
                "knowledge_name": "线性表的顺序存储与链式存储",
                "chapter_name": "数据结构",
                "core_definition": (
                    "核心理解：顺序表连续存储，链表使用指针连接结点。\n\n"
                    "常见考法：比较时间复杂度。\n\n"
                    "易错提醒：链表找到位置通常需要 O(n)。\n\n"
                    "掌握标准：能够根据操作频率选择存储结构。"
                ),
                "keywords_json": '["顺序表", "链表"]',
                "exam_question_styles_json": '["比较时间复杂度"]',
                "related_concepts_json": '["栈与队列"]',
                "pitfalls_json": '["不要忽略查找位置的时间"]',
                "source_text": "这段来源内容不应出现在学生打印版。",
                "source_page": "第2页",
                "material_filename": "学校考纲.pdf",
                "is_ai_expansion": 1,
                "uncertainty_note": "AI 内容待审校。",
            },
            {
                "knowledge_name": "进程与线程",
                "chapter_name": "操作系统",
                "core_definition": "进程是资源分配的基本单位，线程是调度的基本单位。",
                "keywords_json": '["进程", "线程"]',
            },
        ]

        docx_bytes = build_knowledge_memorization_docx(points, subject="408综合")

        self.assertTrue(docx_bytes.startswith(b"PK"))
        document = Document(io.BytesIO(docx_bytes))
        text = _document_text(document)
        self.assertIn("408综合背诵手册", text)
        self.assertIn("线性表的顺序存储与链式存储", text)
        self.assertIn("进程与线程", text)
        self.assertIn("背诵自测", text)
        self.assertEqual(text.count("常见考法："), 1)
        for forbidden in (
            "来源内容",
            "第2页",
            "学校考纲.pdf",
            "AI 内容待审校",
            "审校说明",
            "生成时间",
            "修改意见",
        ):
            self.assertNotIn(forbidden, text)
        section = document.sections[0]
        self.assertAlmostEqual(section.page_width / 360000, 21.0, places=1)
        self.assertAlmostEqual(section.page_height / 360000, 29.7, places=1)

    def test_rejects_empty_document(self):
        with self.assertRaisesRegex(ValueError, "没有可导出"):
            build_knowledge_memorization_docx([], subject="专业课")


if __name__ == "__main__":
    unittest.main()
