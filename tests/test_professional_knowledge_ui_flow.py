import os
import shutil
import sqlite3
import tempfile
import time
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document
from streamlit.testing.v1 import AppTest


PROJECT_DIR = Path(__file__).resolve().parents[1]


def _find_by_label(elements, label, occurrence=0):
    matches = [element for element in elements if element.label == label]
    if len(matches) <= occurrence:
        raise AssertionError(f"未找到控件：{label}（序号 {occurrence}）")
    return matches[occurrence]


class ProfessionalKnowledgeUiFlowTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = Path(tempfile.mkdtemp())
        self.db_path = self.temp_dir / "memory.db"
        self.original_memory_db_env = os.environ.get("MEMORY_DB")
        self.original_api_key_env = os.environ.get("AI_API_KEY")
        self.original_standalone_user_env = os.environ.get("KAOYAN_STANDALONE_USER_ID")
        os.environ["MEMORY_DB"] = str(self.db_path)
        os.environ["AI_API_KEY"] = ""
        os.environ["KAOYAN_STANDALONE_USER_ID"] = "1"

        import knowledge_base
        import professional_knowledge.catalog as catalog
        import repositories.knowledge_repo as knowledge_repo
        import repositories.wrong_question_repo as wrong_question_repo
        import services.professional_knowledge_task_service as task_service

        self.knowledge_base = knowledge_base
        self.catalog = catalog
        self.knowledge_repo = knowledge_repo
        self.wrong_question_repo = wrong_question_repo
        self.task_service = task_service
        self.original_memory_db = knowledge_base.MEMORY_DB
        self.original_wrong_question_db = wrong_question_repo.MEMORY_DB
        self.original_custom_config = catalog.CUSTOM_SUBJECTS_CONFIG_PATH
        self.original_tasks_dir = task_service.TASKS_DIR
        self.original_llm_call = knowledge_base._call_llm_api

        knowledge_base.MEMORY_DB = str(self.db_path)
        wrong_question_repo.MEMORY_DB = str(self.db_path)
        catalog.CUSTOM_SUBJECTS_CONFIG_PATH = self.temp_dir / "custom_subjects.json"
        task_service.TASKS_DIR = self.temp_dir / "tasks"

    def tearDown(self):
        self.knowledge_base.MEMORY_DB = self.original_memory_db
        self.wrong_question_repo.MEMORY_DB = self.original_wrong_question_db
        self.catalog.CUSTOM_SUBJECTS_CONFIG_PATH = self.original_custom_config
        self.task_service.TASKS_DIR = self.original_tasks_dir
        self.knowledge_base._call_llm_api = self.original_llm_call
        if self.original_memory_db_env is None:
            os.environ.pop("MEMORY_DB", None)
        else:
            os.environ["MEMORY_DB"] = self.original_memory_db_env
        if self.original_api_key_env is None:
            os.environ.pop("AI_API_KEY", None)
        else:
            os.environ["AI_API_KEY"] = self.original_api_key_env
        if self.original_standalone_user_env is None:
            os.environ.pop("KAOYAN_STANDALONE_USER_ID", None)
        else:
            os.environ["KAOYAN_STANDALONE_USER_ID"] = self.original_standalone_user_env
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def _run_app(self):
        app = AppTest.from_file(str(PROJECT_DIR / "app_kb.py"), default_timeout=30).run()
        if app.exception:
            raise AssertionError(app.exception)
        return app

    def _seed_knowledge(self, subject, name):
        with sqlite3.connect(self.db_path) as conn:
            self.knowledge_repo.ensure_knowledge_schema(conn)
            self.knowledge_repo.save_confirmed_knowledge_points(
                conn,
                user_id=1,
                points=[
                    {
                        "knowledge_name": name,
                        "core_definition": f"{name}的核心定义。",
                        "source_text": f"原文依据：{name}",
                    }
                ],
                material_meta={"subject": subject},
                strict=False,
            )
            conn.commit()

    def _enable_fake_professional_ai(self):
        os.environ["AI_API_KEY"] = "test-key"
        calls = {"question": 0, "grade": 0}

        def fake_llm(prompt, **kwargs):
            if "学生回答" in prompt or "阅卷" in prompt:
                calls["grade"] += 1
                return (
                    '{"score":82,"feedback":"要点基本正确，继续补充边界条件。",'
                    '"rating":"good","missed_points":["边界条件"],'
                    '"mistake_reason":"过程说明还可以更完整。",'
                    '"next_review":"复习对应知识点的适用条件。",'
                    '"similar_question":"换一组条件再做一题。"}'
                )
            calls["question"] += 1
            suffix = calls["question"]
            if "单选题" in prompt or "choice" in prompt:
                return (
                    '{"question_type":"choice",'
                    f'"question":"某 AOV 网边集为 A->C、B->C、C->D，第 {suffix} 次生成。下列说法哪一项正确？",'
                    '"options":["A. C 可以排在 A 前面","B. A 和 B 都必须排在 C 前面","C. D 必须排在 C 前面","D. 图中一定存在环"],'
                    '"correct_answer":"B","reference_answer":"A、B 均为 C 的前驱，C 又先于 D。",'
                    '"grading_points":["入度约束","拓扑序列","有向无环图"],'
                    '"similar_question":"换一组先修关系再判断。"}'
                )
            if "填空题" in prompt or "blank" in prompt:
                return (
                    '{"question_type":"blank",'
                    f'"question":"填空：对 AOV 网做拓扑排序时，第 {suffix} 次生成，应优先选择入度为 ______ 的顶点。",'
                    '"options":[],"correct_answer":"0",'
                    '"reference_answer":"填 0。入度为 0 表示所有前驱活动已经完成。",'
                    '"grading_points":["入度为0","前驱活动完成"],'
                    '"similar_question":"换一张图写出第一轮可选顶点。"}'
                )
            if "概念辨析题" in prompt or "concept" in prompt:
                return (
                    '{"question_type":"concept",'
                    f'"question":"请用自己的话解释 AOV 网与拓扑排序的关系，并说明第 {suffix} 次生成时如何判断有向环。",'
                    '"options":[],"correct_answer":"",'
                    '"reference_answer":"AOV 网用顶点表示活动、边表示先后关系；拓扑排序通过反复选择入度为 0 的顶点判断是否存在环。",'
                    '"grading_points":["活动顶点","先后关系","入度为0","有向环"],'
                    '"similar_question":"解释 AOE 网和 AOV 网的区别。"}'
                )
            return (
                '{"question_type":"application",'
                f'"question":"已知活动依赖 A->C、B->C、C->D，第 {suffix} 次生成。请给出一种拓扑序列并说明判断依据。",'
                '"options":[],"correct_answer":"",'
                '"reference_answer":"A、B 必须在 C 前，C 必须在 D 前，因此 A,B,C,D 或 B,A,C,D 均可。",'
                '"grading_points":["前驱关系","拓扑序列","有向无环图"],'
                '"similar_question":"把边改为 A->B、C->B，再给出拓扑序列。"}'
            )

        self.knowledge_base._call_llm_api = fake_llm
        return calls

    @staticmethod
    def _markdown_contains(app, text):
        return any(text in str(element.value) for element in app.markdown)

    @staticmethod
    def _info_contains(app, text):
        return any(text in str(element.value) for element in app.info)

    def test_formal_page_keeps_knowledge_base_and_removes_legacy_tools(self):
        self._seed_knowledge("408综合", "栈")
        app = self._run_app()
        _find_by_label(app.radio, "专业课功能").set_value("高级知识条目管理").run()
        if app.exception:
            raise AssertionError(app.exception)

        self.assertTrue(self._markdown_contains(app, "专业课学习"))
        self.assertTrue(self._markdown_contains(app, "408综合知识库"))
        self.assertTrue(self._markdown_contains(app, "栈"))
        self.assertTrue(any(item.label == "上传考试大纲" for item in app.get("file_uploader")))
        self.assertFalse(any(item.label == "粘贴文本" for item in app.text_area))
        self.assertFalse(any(item.label == "确认文本并开始识别" for item in app.button))
        self.assertFalse(any(item.label == "继续处理这份资料" for item in app.button))
        self.assertFalse(self._markdown_contains(app, "高级校对与知识库工具"))
        self.assertFalse(self._markdown_contains(app, "默认全部参与回答"))
        self.assertFalse(self._markdown_contains(app, "左侧负责筛选"))
        self.assertFalse(self._markdown_contains(app, "广告与个人履历不会入库"))

    def test_408_builtin_library_is_available_without_uploads(self):
        app = self._run_app()

        self.assertTrue(self._markdown_contains(app, "408综合知识库"))
        self.assertTrue(any(item.label == "考试科目" for item in app.selectbox))
        self.assertFalse(any(item.label == "资料范围" for item in app.selectbox))
        download_buttons = [
            item
            for item in app.get("download_button")
            if item.label == "下载背诵版 DOCX"
        ]
        self.assertEqual(len(download_buttons), 1)
        self.assertTrue(download_buttons[0].proto.url.endswith(".docx"))
        with sqlite3.connect(self.db_path) as conn:
            count = conn.execute(
                "SELECT COUNT(*) FROM user_knowledge WHERE source_type='builtin_408'"
            ).fetchone()[0]
            has_linear_list = conn.execute(
                """SELECT COUNT(*)
                   FROM user_knowledge
                   WHERE source_type='builtin_408'
                     AND knowledge_name='线性表的顺序存储与链式存储'"""
            ).fetchone()[0]
            chapters = {
                row[0]
                for row in conn.execute(
                    "SELECT DISTINCT chapter_name FROM user_knowledge WHERE source_type='builtin_408'"
                ).fetchall()
            }
        self.assertGreaterEqual(count, 120)
        self.assertEqual(has_linear_list, 1)
        self.assertTrue({"数据结构", "计算机组成原理", "操作系统", "计算机网络"}.issubset(chapters))

    def test_knowledge_library_can_filter_school_priority_points(self):
        from repositories.professional_syllabus_repo import (
            create_syllabus_analysis,
            save_syllabus_analysis_result,
        )

        with sqlite3.connect(self.db_path) as conn:
            analysis = create_syllabus_analysis(
                conn,
                user_id=1,
                subject="408综合",
                source_ids=[],
            )
            save_syllabus_analysis_result(
                conn,
                analysis["id"],
                school_focus=[
                    {
                        "exam_subject": "数据结构",
                        "intensity": "高",
                        "summary": "目标院校强调 AOV 网和拓扑排序。",
                        "evidence": ["AOV 网与活动排序"],
                    }
                ],
                priority_points=[
                    {
                        "knowledge_name": "AOV 网与活动排序",
                        "exam_subject": "数据结构",
                        "reason": "学校考纲直接命中。",
                    }
                ],
                phase_plan=[],
                raw_summary="先抓拓扑相关内容。",
            )
            conn.commit()

        app = self._run_app()
        _find_by_label(app.selectbox, "学习范围").set_value("学校优先").run()
        if app.exception:
            raise AssertionError(app.exception)

        self.assertTrue(self._markdown_contains(app, "AOV 网与活动排序"))
        self.assertFalse(self._markdown_contains(app, "B 树和 B+ 树是多路平衡查找树"))

    def test_syllabus_sources_are_detected_separately(self):
        self.assertTrue(self.knowledge_base._is_syllabus_source({"chapter_name": "学校考纲 - 408大纲.pdf"}))
        self.assertFalse(self.knowledge_base._is_syllabus_source({"chapter_name": "数据结构讲义"}))

    def test_408_builtin_chat_answers_without_uploaded_sources(self):
        answer = self.knowledge_base._answer_subject_question(
            1,
            "408综合",
            [],
            "TCP 三次握手为什么不是两次？",
            answer_mode="custom",
        )

        self.assertIn("TCP", answer)
        self.assertIn("三次握手", answer)
        self.assertNotIn("请先在左侧至少勾选一份资料", answer)

    def test_multiple_knowledge_rows_can_stay_expanded(self):
        app = self._run_app()

        _find_by_label(app.button, "展开查看", occurrence=0).click().run()
        if app.exception:
            raise AssertionError(app.exception)
        _find_by_label(app.button, "展开查看", occurrence=0).click().run()
        if app.exception:
            raise AssertionError(app.exception)

        self.assertTrue(self._markdown_contains(app, "AOV 网用顶点表示活动"))
        self.assertTrue(self._markdown_contains(app, "B 树和 B+ 树是多路平衡查找树"))
        self.assertTrue(self._markdown_contains(app, "掌握标准"))
        self.assertFalse(self._markdown_contains(app, "真题风格例题"))

    def test_quiz_panel_appears_under_selected_knowledge(self):
        self._enable_fake_professional_ai()
        app = self._run_app()

        _find_by_label(app.button, "出题", occurrence=0).click().run()
        if app.exception:
            raise AssertionError(app.exception)

        self.assertTrue(self._markdown_contains(app, "AOV 网与活动排序 · 综合应用题"))
        self.assertEqual(len(app.info), 1)
        self.assertEqual(len(app.error), 0)
        self.assertFalse(self._markdown_contains(app, ">题目<"))
        self.assertTrue(any(item.label == "换一道题" for item in app.button))
        _find_by_label(app.button, "保存本题").click().run()
        if app.exception:
            raise AssertionError(app.exception)
        with sqlite3.connect(self.db_path) as conn:
            saved_count = conn.execute("SELECT COUNT(*) FROM professional_saved_questions").fetchone()[0]
        self.assertEqual(saved_count, 1)
        self.assertTrue(self._markdown_contains(app, "你的回答") or any(item.label == "你的回答" for item in app.text_area))
        workspace_nav = _find_by_label(app.radio, "专业课功能")
        self.assertEqual(workspace_nav.value, "知识库")

    def test_regenerate_and_save_stays_in_knowledge_view_and_keeps_multiple_questions(self):
        self._enable_fake_professional_ai()
        app = self._run_app()

        _find_by_label(app.button, "出题", occurrence=0).click().run()
        if app.exception:
            raise AssertionError(app.exception)
        _find_by_label(app.button, "保存本题").click().run()
        if app.exception:
            raise AssertionError(app.exception)
        _find_by_label(app.button, "换一道题").click().run()
        if app.exception:
            raise AssertionError(app.exception)
        _find_by_label(app.button, "保存本题").click().run()
        if app.exception:
            raise AssertionError(app.exception)

        workspace_nav = _find_by_label(app.radio, "专业课功能")
        self.assertEqual(workspace_nav.value, "知识库")
        with sqlite3.connect(self.db_path) as conn:
            saved_count = conn.execute("SELECT COUNT(*) FROM professional_saved_questions").fetchone()[0]
        self.assertEqual(saved_count, 2)

    def test_choice_question_has_options_and_can_be_saved(self):
        self._enable_fake_professional_ai()
        app = self._run_app()

        _find_by_label(app.button, "选择题", occurrence=0).click().run()
        if app.exception:
            raise AssertionError(app.exception)

        self.assertTrue(self._markdown_contains(app, "AOV 网与活动排序 · 选择题"))
        self.assertTrue(self._markdown_contains(app, "A."))
        self.assertTrue(self._markdown_contains(app, "B."))
        _find_by_label(app.button, "保存本题").click().run()
        if app.exception:
            raise AssertionError(app.exception)
        with sqlite3.connect(self.db_path) as conn:
            row = conn.execute(
                "SELECT question, reference_answer, source_mode FROM professional_saved_questions"
            ).fetchone()
        self.assertIn("A.", row[0])
        self.assertIn("正确答案", row[1])
        self.assertEqual(row[2], "choice")

    def test_local_grading_returns_actionable_review_fields(self):
        point = {
            "knowledge_name": "Cache 映射与替换",
            "core_definition": "Cache 映射包括直接映射、全相联和组相联。",
            "keywords_json": '["Cache","直接映射","组相联","LRU"]',
        }

        result = self.knowledge_base._grade_professional_answer(
            point,
            "某机采用组相联 Cache，请说明地址字段划分。",
            "我只记得和缓存有关。",
            "需要说明标记、组号和块内地址。",
            ["标记", "组号", "块内地址"],
            "application",
        )

        self.assertIn("missed_points", result)
        self.assertIn("mistake_reason", result)
        self.assertIn("next_review", result)
        self.assertIn("similar_question", result)
        self.assertEqual(result["grading_source"], "local_estimate")
        self.assertFalse(result["is_authoritative"])

    def test_fast_grading_does_not_call_llm_even_when_key_exists(self):
        os.environ["AI_API_KEY"] = "test-key"
        self.knowledge_base._call_llm_api = lambda *args, **kwargs: (_ for _ in ()).throw(
            AssertionError("fast grading should not call LLM")
        )

        result = self.knowledge_base._grade_professional_answer(
            {"knowledge_name": "Cache", "core_definition": "Cache 利用局部性。"},
            "解释 Cache 的作用。",
            "利用局部性加快访问。",
            "利用局部性减少平均访存时间。",
            ["局部性", "平均访存时间"],
            "application",
            use_ai=False,
        )

        self.assertGreaterEqual(result["score"], 0)

    def test_ai_grading_prompt_requires_score_breakdown_and_standard_answer(self):
        os.environ["AI_API_KEY"] = "test-key"
        captured = {}

        def fake_llm(prompt, **kwargs):
            captured["prompt"] = prompt
            captured["max_tokens"] = kwargs.get("max_tokens")
            return (
                '{"score":76,"rating":"good","feedback":"过程基本对，但缺少复杂度分析。",'
                '"standard_answer":"应先写地址字段划分，再说明查找流程。",'
                '"score_breakdown":[{"point":"字段划分","score":30,"max_score":40,"comment":"写出了组号和块内地址，但标记位说明不完整。"}],'
                '"missed_points":["标记位位数","复杂度/代价分析"],'
                '"mistake_reason":"只写结论，没有把地址位数和组数对应起来。",'
                '"corrected_answer":"地址分为标记、组号、块内地址三部分，并按组号定位。",'
                '"next_review":"复习 Cache 组相联映射的地址字段划分。",'
                '"similar_question":"给定 Cache 参数，重新划分主存地址字段。"}'
            )

        self.knowledge_base._call_llm_api = fake_llm

        result = self.knowledge_base._grade_professional_answer(
            {"knowledge_name": "Cache 映射", "core_definition": "组相联 Cache 通过组号定位集合。"},
            "某机采用组相联 Cache，请说明地址字段划分。",
            "地址里有组号和偏移。",
            "需要说明标记、组号和块内地址。",
            ["标记", "组号", "块内地址"],
            "application",
            use_ai=True,
        )

        self.assertIn("得分点", captured["prompt"])
        self.assertIn("标准答案", captured["prompt"])
        self.assertIn("学生回答", captured["prompt"])
        self.assertGreaterEqual(captured["max_tokens"], 1200)
        self.assertEqual(result["score"], 76)
        self.assertEqual(result["standard_answer"], "应先写地址字段划分，再说明查找流程。")
        self.assertEqual(result["corrected_answer"], "地址分为标记、组号、块内地址三部分，并按组号定位。")
        self.assertEqual(result["score_breakdown"][0]["point"], "字段划分")
        self.assertEqual(result["grading_source"], "ai")
        self.assertTrue(result["is_authoritative"])

    def test_ai_grading_failure_is_marked_as_local_estimate(self):
        os.environ["AI_API_KEY"] = "test-key"
        self.knowledge_base._call_llm_api = lambda *args, **kwargs: (_ for _ in ()).throw(
            TimeoutError("simulated timeout")
        )

        result = self.knowledge_base._grade_professional_answer(
            {"knowledge_name": "Cache", "core_definition": "Cache 利用局部性。"},
            "解释 Cache 的作用。",
            "利用局部性加快访问。",
            "利用局部性减少平均访存时间。",
            ["局部性", "平均访存时间"],
            "application",
            use_ai=True,
        )

        self.assertEqual(result["grading_source"], "local_estimate")
        self.assertFalse(result["is_authoritative"])
        self.assertIn("不计入正式学习进度", result["grading_warning"])

    def test_llm_api_uses_configured_default_model(self):
        captured = {}
        original_completion = self.knowledge_base.simple_prompt_completion

        def fake_completion(prompt, **kwargs):
            captured.update(kwargs)
            return "OK"

        self.knowledge_base.simple_prompt_completion = fake_completion
        try:
            result = self.knowledge_base._call_llm_api("ping", max_tokens=32)
        finally:
            self.knowledge_base.simple_prompt_completion = original_completion

        self.assertEqual(result, "OK")
        self.assertIsNone(captured.get("model"))
        self.assertEqual(captured.get("max_tokens"), 32)

    def test_ai_question_generation_accepts_chinese_json_fields(self):
        os.environ["AI_API_KEY"] = "test-key"
        self.knowledge_base._call_llm_api = lambda *args, **kwargs: (
            '{"题目":"已知活动依赖为 A->B、A->C、B->D、C->D，请给出一种拓扑序列并说明判断依据。",'
            '"参考答案":"A 必须在 B、C 之前，B、C 必须在 D 之前，因此 A,B,C,D 或 A,C,B,D 均可。",'
            '"评分点":["入度为0","依赖关系","有向无环图"]}'
        )
        point = {
            "knowledge_name": "AOV 网与活动排序",
            "core_definition": "AOV 网用于表达活动之间的先后关系。",
            "keywords_json": '["AOV网","拓扑序列","入度"]',
        }

        generated, warning = self.knowledge_base._generate_professional_question_with_ai(point, "quiz")

        self.assertEqual(warning, "")
        self.assertIn("A->B", generated["question"])
        self.assertIn("入度为0", generated["grading_points"])

    def test_primary_question_generation_uses_ai_when_configured(self):
        os.environ["AI_API_KEY"] = "test-key"
        calls = []

        def fake_llm(prompt, **kwargs):
            calls.append(prompt)
            return (
                '{"question_type":"choice",'
                '"question":"某课程先修关系为 A->C、B->C、C->D。若用拓扑排序安排学习顺序，下列说法哪一项正确？",'
                '"options":["A. C 可以排在 A 前面","B. A 和 B 都必须排在 C 前面","C. D 必须排在 C 前面","D. 图中一定存在环"],'
                '"correct_answer":"B",'
                '"reference_answer":"A、B 的入度约束要求二者均在 C 之前，C 再先于 D。",'
                '"grading_points":["入度约束","拓扑序列","有向无环图"],'
                '"similar_question":"把边改为 A->B、A->D、C->D，再判断合法拓扑序列。"}'
            )

        self.knowledge_base._call_llm_api = fake_llm
        point = {
            "knowledge_name": "AOV 网与活动排序",
            "core_definition": "AOV 网用于表达活动之间的先后关系。",
            "keywords_json": '["AOV网","拓扑排序","入度"]',
        }

        generated = self.knowledge_base._generate_professional_question(point, "choice", variant=3)

        self.assertEqual(len(calls), 1)
        self.assertIn("某课程先修关系", generated["question"])
        self.assertEqual(generated["correct_answer"], "B")
        self.assertEqual(len(generated["options"]), 4)

    def test_configured_ai_question_failure_does_not_use_local_template(self):
        os.environ["AI_API_KEY"] = "test-key"

        def broken_llm(prompt, **kwargs):
            raise TimeoutError("simulated timeout")

        self.knowledge_base._call_llm_api = broken_llm
        point = {
            "knowledge_name": "哈夫曼树与哈夫曼编码",
            "core_definition": "哈夫曼树用贪心思想构造最小 WPL 的二叉树。",
            "keywords_json": '["哈夫曼树","WPL","前缀编码"]',
        }

        generated = self.knowledge_base._generate_professional_question(point, "choice", variant=2)

        self.assertTrue(generated.get("generation_failed"))
        self.assertIn("AI 出题暂时不可用", generated["question"])
        self.assertNotIn("围绕“哈夫曼树与哈夫曼编码”做题时", generated["question"])
        self.assertNotIn("若题干改变约束条件", generated["question"])

    def test_successful_regeneration_clears_sticky_failure_state(self):
        active = {
            "question": "AI 出题暂时不可用",
            "generation_failed": True,
            "generation_warning": "超时",
            "result": {"score": 0},
        }
        generated = {
            "question": "给定一个 AOV 网，请写出一种拓扑序列。",
            "reference_answer": "反复选择入度为 0 的顶点。",
            "generation_failed": False,
            "generation_warning": "",
        }

        merged = self.knowledge_base._merge_generated_question_state(active, generated)

        self.assertNotIn("generation_failed", merged)
        self.assertNotIn("generation_warning", merged)
        self.assertNotIn("result", merged)
        self.assertIn("拓扑序列", merged["question"])

    def test_main_question_request_uses_one_transient_retry(self):
        os.environ["AI_API_KEY"] = "test-key"
        captured = {}

        def fake_llm(prompt, **kwargs):
            captured.update(kwargs)
            return (
                '{"question_type":"application",'
                '"question":"给定 AOV 网边集 A->C、B->C、C->D，请写出一种拓扑序列并说明判环方法。",'
                '"options":[],"correct_answer":"",'
                '"reference_answer":"反复输出入度为 0 的顶点；若无法输出全部顶点则存在环。",'
                '"grading_points":["入度更新","拓扑序列","判环"]}'
            )

        self.knowledge_base._call_llm_api = fake_llm
        generated, warning = self.knowledge_base._generate_professional_question_with_ai(
            {
                "knowledge_name": "AOV 网与活动排序",
                "core_definition": "拓扑排序用于有向无环图。",
                "keywords_json": '["AOV网","拓扑排序","入度"]',
            },
            "application",
            variant=2,
        )

        self.assertEqual(warning, "")
        self.assertFalse(generated["generation_failed"])
        self.assertEqual(captured["retries"], 1)

    def test_configured_ai_incomplete_question_does_not_use_local_template(self):
        os.environ["AI_API_KEY"] = "test-key"
        self.knowledge_base._call_llm_api = lambda *args, **kwargs: '{"question":"题目","reference_answer":""}'
        point = {
            "knowledge_name": "哈夫曼树与哈夫曼编码",
            "core_definition": "哈夫曼编码是前缀编码。",
            "keywords_json": '["哈夫曼树","WPL","前缀编码"]',
        }

        generated = self.knowledge_base._generate_professional_question(point, "blank", variant=3)

        self.assertTrue(generated.get("generation_failed"))
        self.assertIn("AI 返回的题干或参考答案不完整", generated.get("generation_warning", ""))
        self.assertNotIn("填空：做“哈夫曼树与哈夫曼编码”相关题时", generated["question"])

    def test_truncated_ai_question_json_can_still_be_used(self):
        os.environ["AI_API_KEY"] = "test-key"
        self.knowledge_base._call_llm_api = lambda *args, **kwargs: (
            '```json\n{"question_type":"application",'
            '"question":"给定字符 A,B,C,D 的权值分别为 2,5,7,9，请构造哈夫曼树并计算 WPL。",'
            '"options":[],'
            '"correct_answer":"",'
            '"reference_answer":"先选 2 和 5 合并为 7，再按最小权值继续合并，最后计算每个叶子的带权路径长度之和。"'
        )
        point = {
            "knowledge_name": "哈夫曼树与哈夫曼编码",
            "core_definition": "哈夫曼树是带权路径长度最小的二叉树。",
            "keywords_json": '["哈夫曼树","WPL","前缀编码"]',
        }

        generated = self.knowledge_base._generate_professional_question(point, "application", variant=4)

        self.assertFalse(generated.get("generation_failed"))
        self.assertIn("权值", generated["question"])
        self.assertIn("带权路径长度", generated["reference_answer"])

    def test_ai_question_with_only_question_completes_reference(self):
        os.environ["AI_API_KEY"] = "test-key"
        calls = {"count": 0}

        def fake_llm(*args, **kwargs):
            calls["count"] += 1
            if calls["count"] == 1:
                return (
                    '{"question":"给定字符 A,B,C,D 的权值分别为 2,5,7,9，'
                    '请构造哈夫曼树并计算 WPL。"'
                )
            return '{"reference_answer":"按最小权值反复合并，并计算各叶子权值乘路径长度之和。","grading_points":["合并顺序","WPL计算"]}'

        self.knowledge_base._call_llm_api = fake_llm
        point = {
            "knowledge_name": "哈夫曼树与哈夫曼编码",
            "core_definition": "哈夫曼树是带权路径长度最小的二叉树。",
            "keywords_json": '["哈夫曼树","WPL","前缀编码"]',
        }

        generated = self.knowledge_base._generate_professional_question(point, "application", variant=5)

        self.assertFalse(generated.get("generation_failed"))
        self.assertEqual(calls["count"], 2)
        self.assertIn("合并", generated["reference_answer"])
        self.assertEqual(generated["grading_points"], ["合并顺序", "WPL计算"])

    def test_clean_assistant_answer_removes_meta_text_and_source_markers(self):
        raw = (
            "现在开始写。\n"
            "注意：用户问题要求不要因为篇幅只写前半部分。\n\n"
            "数据结构\n"
            "- 线性表是顺序存储和链式存储的基础。[来源1，第2页]\n"
            "- 栈和队列常用于过程模拟。[来源1]\n"
        )

        cleaned = self.knowledge_base._clean_assistant_answer(raw)

        self.assertNotIn("现在开始写", cleaned)
        self.assertNotIn("用户问题", cleaned)
        self.assertNotIn("来源1", cleaned)
        self.assertIn("数据结构", cleaned)
        self.assertIn("线性表", cleaned)

    def test_concept_self_test_has_real_question_and_no_blocking_spinner(self):
        self._enable_fake_professional_ai()
        app = self._run_app()

        _find_by_label(app.button, "概念自测", occurrence=0).click().run()
        if app.exception:
            raise AssertionError(app.exception)

        self.assertTrue(self._markdown_contains(app, "AOV 网与活动排序 · 概念自测"))
        self.assertEqual(len(app.info), 1)
        self.assertEqual(len(app.error), 0)
        self.assertFalse(self._markdown_contains(app, "正在准备概念自测"))
        self.assertFalse(self._markdown_contains(app, "正在按这条知识点出题"))

    def test_blank_question_mode_is_available(self):
        self._enable_fake_professional_ai()
        app = self._run_app()

        _find_by_label(app.button, "填空题", occurrence=0).click().run()
        if app.exception:
            raise AssertionError(app.exception)

        self.assertTrue(self._markdown_contains(app, "填空题"))
        self.assertEqual(len(app.info), 1)
        self.assertEqual(len(app.error), 0)

    def test_blank_answer_array_is_rendered_as_readable_text(self):
        normalized = self.knowledge_base._normalize_professional_question_payload(
            {
                "question_type": "blank",
                "question": "快速排序平均时间复杂度为 ______，归并排序是 ______ 排序。",
                "options": [],
                "correct_answer": ["O(nlogn)", "稳定"],
                "reference_answer": "第一空填 O(nlogn)，第二空填稳定。",
                "grading_points": ["复杂度", "稳定性"],
            },
            {"knowledge_name": "内部排序算法比较"},
            "blank",
            self.knowledge_base._empty_professional_question_payload("blank"),
        )

        self.assertEqual(normalized["correct_answer"], "O(nlogn)；稳定")
        self.assertTrue(
            self.knowledge_base._is_valid_professional_question_for_point(
                normalized,
                "blank",
                {"knowledge_name": "内部排序算法比较"},
            )
        )

    def test_conflicting_ai_reference_is_rejected(self):
        generated = {
            "question_type": "blank",
            "question": "快速排序最坏时间复杂度为 ______。",
            "options": [],
            "correct_answer": "O(n)",
            "reference_answer": "严格来说，快速排序最坏时间复杂度为 O(n^2)，但答案给出 O(n)。",
            "grading_points": ["复杂度"],
        }

        self.assertFalse(
            self.knowledge_base._is_valid_professional_question_for_point(
                generated,
                "blank",
                {"knowledge_name": "内部排序算法比较"},
            )
        )

    def test_blank_question_accepts_supported_shortest_path_answers(self):
        generated = {
            "question_type": "blank",
            "question": "带权有向图从 A 到 D 的边有 A->B=2、A->C=6、B->D=3、C->D=1。先确定的中间顶点是 ______，A 到 D 的最短路径长度为 ______。",
            "options": [],
            "correct_answer": "B；5",
            "reference_answer": "第一空填 B，因为 Dijkstra 初始距离中 B 的距离 2 最小，应先确定 B。第二空填 5，经过 B 松弛后 A 到 D 的距离为 2+3=5，小于经 C 的 7。",
            "grading_points": ["先确定最小距离顶点", "完成松弛", "写出最短路径长度"],
        }

        self.assertTrue(
            self.knowledge_base._is_valid_professional_question_for_point(
                generated,
                "blank",
                {"knowledge_name": "最短路径算法", "subject": "数据结构"},
            )
        )

    def test_blank_question_rejects_answer_not_supported_by_reference(self):
        generated = {
            "question_type": "blank",
            "question": "带权有向图从 A 到 D 的边有 A->B=2、A->C=6、B->D=3、C->D=1。先确定的中间顶点是 ______，A 到 D 的最短路径长度为 ______。",
            "options": [],
            "correct_answer": "B；7",
            "reference_answer": "第一空填 B，因为 Dijkstra 初始距离中 B 的距离 2 最小，应先确定 B。第二空填 5，经过 B 松弛后 A 到 D 的距离为 2+3=5，小于经 C 的 7。",
            "grading_points": ["先确定最小距离顶点", "完成松弛", "写出最短路径长度"],
        }

        self.assertFalse(
            self.knowledge_base._is_valid_professional_question_for_point(
                generated,
                "blank",
                {"knowledge_name": "最短路径算法", "subject": "数据结构"},
            )
        )

    def test_blank_marker_variants_are_accepted(self):
        generated = {
            "question_type": "blank",
            "question": "在 Dijkstra 算法中，若存在负权边但无负权回路，应改用（ ）算法处理单源最短路径。",
            "options": [],
            "correct_answer": "Bellman-Ford",
            "reference_answer": "空处填 Bellman-Ford。Dijkstra 依赖每次确定的最短距离不会再变小，负权边会破坏这个性质；Bellman-Ford 通过多轮松弛处理含负权边的单源最短路径。",
            "grading_points": ["识别负权条件", "说明 Dijkstra 限制", "给出替代算法"],
        }

        self.assertTrue(
            self.knowledge_base._is_valid_professional_question_for_point(
                generated,
                "blank",
                {"knowledge_name": "最短路径算法", "subject": "数据结构"},
            )
        )

    def test_question_variant_changes_between_starts(self):
        self._enable_fake_professional_ai()
        point = {
            "id": 999,
            "knowledge_name": "AOV 网与活动排序",
            "core_definition": "AOV 网用于表示活动先后关系。",
            "keywords_json": '["AOV网","拓扑排序","入度"]',
        }

        self.knowledge_base.st.session_state.clear()
        self.knowledge_base._start_professional_study(1, "408综合", point, "choice", "library")
        first = self.knowledge_base.st.session_state[
            self.knowledge_base._study_state_key(1, "408综合", "library")
        ]["question"]
        self.knowledge_base._start_professional_study(1, "408综合", point, "choice", "library")
        second = self.knowledge_base.st.session_state[
            self.knowledge_base._study_state_key(1, "408综合", "library")
        ]["question"]

        self.assertNotEqual(first, second)

    def test_memory_system_starts_unreviewed_points_at_zero(self):
        app = self._run_app()
        _find_by_label(app.radio, "专业课功能").set_value("记忆系统").run()
        if app.exception:
            raise AssertionError(app.exception)

        self.assertTrue(self._markdown_contains(app, "初始知识点按 0% 计算"))
        self.assertTrue(self._markdown_contains(app, "已验证掌握度"))
        self.assertFalse(self._markdown_contains(app, "整体掌握度 30%"))

    def test_school_priority_points_can_be_added_to_review_by_name(self):
        self._seed_knowledge("408综合", "目标院校自定义重点")

        added = self.knowledge_base._mark_priority_points_due_today(
            1,
            "408综合",
            [{"knowledge_name": "目标院校自定义重点", "exam_subject": "数据结构"}],
        )

        self.assertEqual(added, 1)
        with sqlite3.connect(self.db_path) as conn:
            row = conn.execute(
                """SELECT next_review FROM professional_memory
                   WHERE user_id=1 AND subject='408综合'
                     AND knowledge_id=(
                         SELECT id FROM user_knowledge
                         WHERE subject='408综合' AND knowledge_name='目标院校自定义重点'
                         ORDER BY id DESC LIMIT 1
                     )"""
            ).fetchone()
        self.assertIsNotNone(row)
        self.assertTrue(row[0])

    def test_assistant_answer_offers_trimmed_pdf_download(self):
        app = AppTest.from_file(str(PROJECT_DIR / "app_kb.py"), default_timeout=30)
        prompt = "请根据已选资料梳理完整的知识框架，并标出各部分之间的关系。"
        app.session_state["pk_active_chat_slot_1_408综合"] = "outline"
        app.session_state["pk_chat_answer_cache_v2_1_408综合"] = {
            "outline|": {
                "prompt": prompt,
                "answer": (
                    "根据您提供的资料，下面给出完整框架。\n\n"
                    "## 数据结构\n\n1. 线性表、栈和队列。"
                ),
            }
        }
        app.run()
        if app.exception:
            raise AssertionError(app.exception)

        self.assertTrue(
            any(item.label == "导出本回答精简版 PDF" for item in app.get("download_button"))
        )

    def test_quick_chat_switches_cached_answers_instead_of_stacking(self):
        original_answer = self.knowledge_base._answer_subject_question
        calls = []

        def fake_answer(*args, **kwargs):
            calls.append(kwargs.get("answer_mode"))
            prompt = args[-1]
            return "框架回答" if "知识框架" in prompt else "复习清单回答"

        self.knowledge_base._answer_subject_question = fake_answer
        app = self._run_app()

        try:
            _find_by_label(app.button, "梳理知识框架").click().run()
            if app.exception:
                raise AssertionError(app.exception)
            self.assertTrue(self._markdown_contains(app, "框架回答"))

            _find_by_label(app.button, "生成复习清单").click().run()
            if app.exception:
                raise AssertionError(app.exception)
            self.assertFalse(self._markdown_contains(app, "框架回答"))
            self.assertTrue(self._markdown_contains(app, "复习清单回答"))

            _find_by_label(app.button, "梳理知识框架").click().run()
            if app.exception:
                raise AssertionError(app.exception)
            self.assertTrue(self._markdown_contains(app, "框架回答"))
            self.assertFalse(self._markdown_contains(app, "复习清单回答"))
            self.assertIn("outline", calls)
            self.assertIn("review", calls)
        finally:
            self.knowledge_base._answer_subject_question = original_answer

    def test_formal_knowledge_base_follows_selected_subject(self):
        self._seed_knowledge("408综合", "栈")
        self.catalog.save_custom_subject_profile(
            {
                "key": "custom_management",
                "catalog": {
                    "title": "管理学原理",
                    "subject_label": "管理学原理",
                    "status": "已启用",
                    "stage": "测试",
                    "summary": "测试用自定义专业课。",
                    "capabilities": ["知识点抽取"],
                    "source_strategy": "测试数据",
                    "notes": "测试不同专业课知识库隔离。",
                    "enabled": True,
                },
                "local_source": None,
                "max_points": 12,
                "exam_subjects": ["管理学原理"],
                "extraction_guidance": "测试。",
            },
            custom_config_path=self.catalog.CUSTOM_SUBJECTS_CONFIG_PATH,
        )
        self._seed_knowledge("管理学原理", "计划职能")
        app = self._run_app()
        _find_by_label(app.radio, "专业课功能").set_value("高级知识条目管理").run()
        if app.exception:
            raise AssertionError(app.exception)
        self.assertTrue(self._markdown_contains(app, "408综合知识库"))
        self.assertTrue(self._markdown_contains(app, "栈"))
        self.assertFalse(self._markdown_contains(app, "计划职能的核心定义"))

        _find_by_label(app.selectbox, "专业课").set_value("管理学原理").run()
        if app.exception:
            raise AssertionError(app.exception)
        _find_by_label(app.radio, "专业课功能").set_value("高级知识条目管理").run()
        if app.exception:
            raise AssertionError(app.exception)
        self.assertTrue(self._markdown_contains(app, "管理学原理知识库"))
        self.assertTrue(self._markdown_contains(app, "计划职能"))
        self.assertFalse(self._markdown_contains(app, "栈的核心定义"))

    def test_selected_knowledge_and_detail_stay_in_sync(self):
        self._seed_knowledge("408综合", "栈")
        self._seed_knowledge("408综合", "队列")
        app = self._run_app()
        _find_by_label(app.radio, "专业课功能").set_value("高级知识条目管理").run()
        if app.exception:
            raise AssertionError(app.exception)
        selector = _find_by_label(app.radio, "知识条目列表")

        self.assertTrue(self._markdown_contains(app, "队列的核心定义"))
        with sqlite3.connect(self.db_path) as conn:
            target_id = str(
                conn.execute(
                    """SELECT id FROM user_knowledge
                       WHERE subject='408综合' AND knowledge_name='栈'"""
                ).fetchone()[0]
            )
        selector.set_value(target_id).run()
        if app.exception:
            raise AssertionError(app.exception)

        self.assertTrue(self._markdown_contains(app, "栈的核心定义"))
        self.assertFalse(self._markdown_contains(app, "队列的核心定义"))

    def test_repository_option_only_shows_knowledge_name(self):
        label = self.knowledge_base._format_repo_option(
            {
                "knowledge_name": "页式虚拟存储器",
                "is_ai_expansion": True,
                "subject": "408",
                "mastery_state": "待复习",
            }
        )
        self.assertEqual(label, "页式虚拟存储器")

    def test_current_knowledge_card_hides_source_badges(self):
        self._seed_knowledge("408综合", "页式虚拟存储器")
        app = self._run_app()
        self.assertFalse(self._markdown_contains(app, "原文整理"))
        self.assertFalse(self._markdown_contains(app, "AI扩展 · 需核对教材"))

    def test_current_knowledge_can_preview_and_save_ai_expansion(self):
        self._seed_knowledge("408综合", "页式虚拟存储器")
        self.knowledge_base._call_llm_api = lambda *args, **kwargs: (
            "## 关联知识点及关系\n\n- TLB：用于加速地址转换。"
        )
        app = self._run_app()
        _find_by_label(app.radio, "专业课功能").set_value("高级知识条目管理").run()
        if app.exception:
            raise AssertionError(app.exception)

        _find_by_label(app.button, "AI 发散当前条目").click().run()
        if app.exception:
            raise AssertionError(app.exception)
        self.assertTrue(self._markdown_contains(app, "AI 发散预览"))
        self.assertTrue(self._markdown_contains(app, "TLB：用于加速地址转换"))

        _find_by_label(app.button, "保存发散内容").click().run()
        if app.exception:
            raise AssertionError(app.exception)
        with sqlite3.connect(self.db_path) as conn:
            saved = conn.execute(
                "SELECT review_content FROM user_knowledge WHERE knowledge_name=?",
                ("页式虚拟存储器",),
            ).fetchone()[0]
        self.assertIn("TLB：用于加速地址转换", saved)
        self.assertTrue(self._markdown_contains(app, "已保存的 AI 发散内容"))

    def test_current_knowledge_can_preview_manual_web_supplement(self):
        self._seed_knowledge("408综合", "Cache 映射与替换")
        original_search = self.knowledge_base.search_web
        original_llm = self.knowledge_base._call_llm_api
        self.knowledge_base.search_web = lambda *args, **kwargs: [
            {
                "title": "Cache 映射资料",
                "url": "https://example.com/cache",
                "snippet": "直接映射、组相联和全相联的比较。",
            }
        ]
        self.knowledge_base._call_llm_api = lambda *args, **kwargs: (
            "联网补充：[网页1] 可重点比较直接映射、组相联和全相联。"
        )
        app = self._run_app()
        _find_by_label(app.radio, "专业课功能").set_value("高级知识条目管理").run()
        if app.exception:
            raise AssertionError(app.exception)

        try:
            _find_by_label(app.button, "联网补充").click().run()
            if app.exception:
                raise AssertionError(app.exception)
            self.assertTrue(self._markdown_contains(app, "联网补充预览"))
            self.assertTrue(self._markdown_contains(app, "直接映射、组相联和全相联"))
        finally:
            self.knowledge_base.search_web = original_search
            self.knowledge_base._call_llm_api = original_llm

    def test_web_supplement_query_follows_current_subject(self):
        query_408 = self.knowledge_base._build_web_supplement_query(
            {
                "subject": "408综合",
                "subject_area": "数据结构",
                "knowledge_name": "最短路径算法",
            }
        )
        query_history = self.knowledge_base._build_web_supplement_query(
            {
                "subject": "历史学基础",
                "subject_area": "中国古代史",
                "knowledge_name": "均田制",
            }
        )

        self.assertIn("408", query_408)
        self.assertIn("最短路径算法", query_408)
        self.assertIn("历史学基础", query_history)
        self.assertIn("均田制", query_history)
        self.assertNotIn("408", query_history)

    def test_failed_chat_job_is_not_cached_as_a_normal_answer(self):
        original_answer = self.knowledge_base._answer_subject_question
        self.knowledge_base._answer_subject_question = (
            lambda *args, **kwargs: (_ for _ in ()).throw(TimeoutError("simulated timeout"))
        )
        cache_id = "failed-chat-test"
        try:
            job_id = self.knowledge_base._start_chat_answer_background(
                1,
                "408综合",
                [],
                "解释拓扑排序",
                "custom",
                cache_id,
            )
            deadline = time.time() + 2
            job = self.knowledge_base._get_chat_job(job_id)
            while job and job.get("status") == "running" and time.time() < deadline:
                time.sleep(0.01)
                job = self.knowledge_base._get_chat_job(job_id)
            answer_cache = {}
            self.knowledge_base._sync_chat_answer_from_job(answer_cache, cache_id, job_id)
        finally:
            self.knowledge_base._answer_subject_question = original_answer

        self.assertEqual(job["status"], "failed")
        self.assertEqual(answer_cache[cache_id]["status"], "failed")
        self.assertNotIn("answer", answer_cache[cache_id])

    def test_knowledge_card_quiz_records_feedback_and_memory(self):
        self._enable_fake_professional_ai()
        self._seed_knowledge("408综合", "页式虚拟存储器")
        app = self._run_app()

        _find_by_label(app.button, "出题").click().run()
        if app.exception:
            raise AssertionError(app.exception)
        self.assertTrue(self._markdown_contains(app, "综合应用题"))

        _find_by_label(app.text_area, "你的回答").set_value(
            "页式虚拟存储器通过页表完成虚拟地址到物理地址的转换。"
        )
        _find_by_label(app.button, "提交并查看反馈").click().run()
        if app.exception:
            raise AssertionError(app.exception)

        self.assertTrue(self._markdown_contains(app, "本次得分"))
        with sqlite3.connect(self.db_path) as conn:
            record_count = conn.execute(
                "SELECT COUNT(*) FROM professional_study_records WHERE study_mode='application'"
            ).fetchone()[0]
            practiced_memory = conn.execute(
                """SELECT COUNT(*)
                   FROM professional_memory m
                   JOIN user_knowledge k ON k.id=m.knowledge_id
                   WHERE m.subject='408综合'
                     AND k.knowledge_name IS NOT NULL"""
            ).fetchone()[0]
            reviewed_count = conn.execute(
                """SELECT COUNT(*)
                   FROM professional_memory
                   WHERE subject='408综合' AND review_count > 0"""
            ).fetchone()[0]
        self.assertEqual(record_count, 1)
        self.assertGreaterEqual(practiced_memory, 120)
        self.assertEqual(reviewed_count, 1)

    def test_custom_subject_wizard_selects_new_subject(self):
        app = self._run_app()
        _find_by_label(app.text_input, "专业课名称").set_value("管理学原理")
        _find_by_label(app.text_input, "科目 1").set_value("管理学原理")
        _find_by_label(app.button, "新增科目").click().run()
        if app.exception:
            raise AssertionError(app.exception)
        _find_by_label(app.text_input, "科目 2").set_value("组织行为学")
        _find_by_label(app.text_input, "考试代码（可选）").set_value("803")
        _find_by_label(app.text_area, "希望系统重点识别什么（可选）").set_value(
            "优先识别理论流派、代表人物和易混点。"
        )
        _find_by_label(app.button, "创建专业课").click().run()
        if app.exception:
            raise AssertionError(app.exception)

        subject_selector = _find_by_label(app.selectbox, "专业课", occurrence=0)
        self.assertIn("管理学原理", subject_selector.options)
        self.assertEqual(subject_selector.value, "管理学原理")
        self.assertTrue(self.catalog.CUSTOM_SUBJECTS_CONFIG_PATH.exists())

    def test_removed_prompt_panels_are_not_rendered(self):
        app = self._run_app()

        self.assertFalse(self._markdown_contains(app, "学校考纲分析"))
        self.assertFalse(any(item.label == "本机资料文件夹（可选，仅本地部署）" for item in app.text_input))

    def test_feynman_uses_ai_grading_when_api_key_is_configured(self):
        os.environ["AI_API_KEY"] = "test-key"
        calls = []
        original_grade = self.knowledge_base._grade_professional_answer

        def fake_grade(*args, **kwargs):
            calls.append(kwargs.get("use_ai"))
            return {
                "score": 82,
                "feedback": "概念基本讲清楚了。",
                "rating": "good",
                "missed_points": ["易混点"],
                "mistake_reason": "对边界条件说明偏少。",
                "next_review": "复习拓扑排序的适用条件。",
                "similar_question": "换一个 AOV 网再讲一遍。",
            }

        self.knowledge_base._grade_professional_answer = fake_grade
        try:
            app = self._run_app()
            _find_by_label(app.radio, "专业课功能").set_value("费曼学习法").run()
            if app.exception:
                raise AssertionError(app.exception)
            _find_by_label(app.text_area, "用自己的话讲一遍").set_value(
                "AOV 网用顶点表示活动，用边表示先后约束，拓扑排序可以判断是否存在环。"
            ).run()
            _find_by_label(app.button, "提交讲解").click().run()
            if app.exception:
                raise AssertionError(app.exception)
        finally:
            self.knowledge_base._grade_professional_answer = original_grade

        self.assertEqual(calls, [True])

    def test_saved_question_review_uses_ai_grading_when_api_key_is_configured(self):
        os.environ["AI_API_KEY"] = "test-key"
        self._seed_knowledge("408综合", "AOV 网与活动排序")
        with sqlite3.connect(self.db_path) as conn:
            point_id = conn.execute(
                "SELECT id FROM user_knowledge WHERE subject=? AND knowledge_name=?",
                ("408综合", "AOV 网与活动排序"),
            ).fetchone()[0]
            from repositories.professional_learning_repo import save_generated_question

            save_generated_question(
                conn,
                user_id=1,
                subject="408综合",
                knowledge_id=point_id,
                question="请给出 AOV 网的一个拓扑序列，并说明判断依据。",
                reference_answer="先选入度为 0 的顶点，若所有顶点都输出则无环。",
                grading_points=["入度为0", "拓扑序列", "有向无环图"],
                source_mode="application",
            )
            conn.commit()

        calls = []
        original_grade = self.knowledge_base._grade_professional_answer

        def fake_grade(*args, **kwargs):
            calls.append(kwargs.get("use_ai"))
            return {
                "score": 86,
                "feedback": "关键步骤基本完整。",
                "rating": "good",
                "missed_points": ["环的判断"],
                "mistake_reason": "最后一步没有说明未输出顶点时表示有环。",
                "next_review": "复习拓扑排序判环过程。",
                "similar_question": "换一张 AOV 网再写拓扑序列。",
            }

        self.knowledge_base._grade_professional_answer = fake_grade
        try:
            app = self._run_app()
            _find_by_label(app.selectbox, "专业课").set_value("408综合").run()
            _find_by_label(app.radio, "专业课功能").set_value("复习挑战").run()
            if app.exception:
                raise AssertionError(app.exception)
            _find_by_label(app.text_area, "复练回答").set_value(
                "先反复选择入度为 0 的顶点输出，输出完成说明不存在环。"
            ).run()
            _find_by_label(app.button, "提交复练并记录").click().run()
            if app.exception:
                raise AssertionError(app.exception)
        finally:
            self.knowledge_base._grade_professional_answer = original_grade

        self.assertEqual(calls, [True])

    def test_configured_subject_can_be_removed_after_confirmation(self):
        app = self._run_app()
        subject_selector = _find_by_label(app.selectbox, "专业课")
        self.assertIn("408综合", subject_selector.options)
        self.assertNotIn("医学考研", subject_selector.options)

        _find_by_label(app.button, "删除专业课").click().run()
        if app.exception:
            raise AssertionError(app.exception)
        _find_by_label(app.checkbox, "我确认移除“408综合”").set_value(True)
        _find_by_label(app.button, "确认删除专业课").click().run()
        if app.exception:
            raise AssertionError(app.exception)

        disabled_profile = self.catalog.get_rag_knowledge_base_by_subject("408综合")
        self.assertIsNotNone(disabled_profile)
        self.assertFalse(disabled_profile.enabled)

    def test_workbench_upload_is_queued_without_running_synchronous_processing(self):
        started_jobs = []
        original_start = self.knowledge_base._start_workbench_upload_background

        class FakeUpload:
            name = "目标院校408考纲.txt"

            def getvalue(self):
                return "数据结构重点考 AOV 网和拓扑排序。".encode("utf-8")

        self.knowledge_base._start_workbench_upload_background = started_jobs.append
        try:
            queued, warnings = self.knowledge_base._queue_workbench_uploads(
                1,
                "408综合",
                [FakeUpload()],
            )
        finally:
            self.knowledge_base._start_workbench_upload_background = original_start

        self.assertEqual(queued, 1)
        self.assertEqual(warnings, [])
        self.assertEqual(len(started_jobs), 1)
        with sqlite3.connect(self.db_path) as conn:
            material = conn.execute(
                "SELECT filename, processing_status FROM user_materials"
            ).fetchone()
        self.assertEqual(material[0], "目标院校408考纲.txt")
        self.assertEqual(material[1], "pending")

    def test_personal_material_upload_is_indexed_and_visible_in_advanced_management(self):
        started_jobs = []
        original_start = self.knowledge_base._start_workbench_upload_background
        original_extract = self.knowledge_base.extract_knowledge_points_as_drafts

        class FakeUpload:
            name = "个人复习笔记.txt"

            def getvalue(self):
                return (
                    "拓扑排序用于有向无环图。每次选择入度为零的顶点输出，"
                    "并删除该顶点及其出边；若最终输出顶点数少于图中顶点数，则图中存在环。"
                ).encode("utf-8")

        from schemas.knowledge_schema import KnowledgePointDraft

        self.knowledge_base._start_workbench_upload_background = started_jobs.append
        self.knowledge_base.extract_knowledge_points_as_drafts = (
            lambda **kwargs: (
                [
                    KnowledgePointDraft(
                        knowledge_name="拓扑排序判环",
                        knowledge_type="个人资料知识点",
                        subject="408综合",
                        chapter_name="个人复习笔记",
                        core_definition="反复输出入度为零的顶点；若不能输出全部顶点，则有向图存在环。",
                        keywords=["拓扑排序", "入度", "有向无环图"],
                        source_text=kwargs["text"],
                    )
                ],
                [],
            )
        )
        try:
            queued, warnings = self.knowledge_base._queue_workbench_uploads(
                1,
                "408综合",
                [FakeUpload()],
            )
            self.assertEqual(queued, 1)
            self.assertEqual(warnings, [])
            self.assertEqual(len(started_jobs), 1)
            self.knowledge_base._process_workbench_upload_job(started_jobs[0])
        finally:
            self.knowledge_base._start_workbench_upload_background = original_start
            self.knowledge_base.extract_knowledge_points_as_drafts = original_extract

        with sqlite3.connect(self.db_path) as conn:
            source_state = conn.execute(
                """SELECT processing_status, knowledge_count
                   FROM user_materials WHERE filename=?""",
                ("个人复习笔记.txt",),
            ).fetchone()
            knowledge_row = conn.execute(
                """SELECT knowledge_name, material_filename
                   FROM user_knowledge WHERE knowledge_name=?""",
                ("拓扑排序判环",),
            ).fetchone()
        self.assertEqual(source_state, ("done", 1))
        self.assertEqual(knowledge_row, ("拓扑排序判环", "个人复习笔记.txt"))

        app = self._run_app()
        _find_by_label(app.text_input, "搜索知识库").set_value("拓扑排序判环").run()
        if app.exception:
            raise AssertionError(app.exception)
        self.assertTrue(self._markdown_contains(app, "拓扑排序判环"))
        self.assertTrue(self._markdown_contains(app, "来自个人资料"))

        app.run()
        if app.exception:
            raise AssertionError(app.exception)
        self.assertTrue(self._markdown_contains(app, "拓扑排序判环"))

        _find_by_label(app.radio, "专业课功能").set_value("高级知识条目管理").run()
        if app.exception:
            raise AssertionError(app.exception)
        self.assertTrue(self._markdown_contains(app, "拓扑排序判环"))

    def test_docx_personal_material_reaches_the_same_knowledge_indexing_pipeline(self):
        started_jobs = []
        original_start = self.knowledge_base._start_workbench_upload_background
        original_extract = self.knowledge_base.extract_knowledge_points_as_drafts
        buffer = BytesIO()
        document = Document()
        document.add_heading("中国近现代史", level=1)
        document.add_paragraph("洋务运动以自强、求富为口号，兴办近代军事工业和民用工业。")
        document.save(buffer)

        class FakeUpload:
            name = "中国近现代史.docx"

            def getvalue(self):
                return buffer.getvalue()

        from schemas.knowledge_schema import KnowledgePointDraft

        self.knowledge_base._start_workbench_upload_background = started_jobs.append
        self.knowledge_base.extract_knowledge_points_as_drafts = (
            lambda **kwargs: (
                [
                    KnowledgePointDraft(
                        knowledge_name="洋务运动",
                        knowledge_type="个人资料知识点",
                        subject="历史学统考",
                        chapter_name="中国近现代史",
                        core_definition="洋务派以自强、求富为口号推进近代化实践。",
                        keywords=["洋务运动", "自强", "求富"],
                        source_text=kwargs["text"],
                    )
                ],
                [],
            )
        )
        try:
            queued, warnings = self.knowledge_base._queue_workbench_uploads(
                1,
                "历史学统考",
                [FakeUpload()],
            )
            self.assertEqual((queued, warnings), (1, []))
            self.knowledge_base._process_workbench_upload_job(started_jobs[0])
        finally:
            self.knowledge_base._start_workbench_upload_background = original_start
            self.knowledge_base.extract_knowledge_points_as_drafts = original_extract

        with sqlite3.connect(self.db_path) as conn:
            source_state = conn.execute(
                """SELECT source_type, process_method, processing_status, knowledge_count
                   FROM user_materials WHERE filename=?""",
                ("中国近现代史.docx",),
            ).fetchone()
            knowledge_row = conn.execute(
                """SELECT knowledge_name, material_filename
                   FROM user_knowledge WHERE knowledge_name=?""",
                ("洋务运动",),
            ).fetchone()
        self.assertEqual(source_state, ("docx", "docx_text_extract", "done", 1))
        self.assertEqual(knowledge_row, ("洋务运动", "中国近现代史.docx"))

    def test_workbench_reprocess_is_queued_without_blocking_page(self):
        started_jobs = []
        original_start = self.knowledge_base._start_workbench_upload_background
        source_path = self.knowledge_base._persist_user_material_file(
            1,
            "old-outline.txt",
            "数据结构重点考拓扑排序。".encode("utf-8"),
        )
        with sqlite3.connect(self.db_path) as conn:
            self.knowledge_repo.ensure_knowledge_schema(conn)
            from repositories.material_repo import create_material

            material = create_material(
                conn,
                user_id=1,
                subject="408综合",
                filename="old-outline.txt",
                chapter_name="旧考纲",
                file_path=str(source_path),
                file_type="txt",
                processing_status="done",
            )
            self.knowledge_repo.save_confirmed_knowledge_points(
                conn,
                user_id=1,
                points=[{"knowledge_name": "旧知识点", "core_definition": "旧内容"}],
                material_meta={"material_id": material["id"], "subject": "408综合"},
            )
            conn.commit()

        source = {
            "id": material["id"],
            "subject": "408综合",
            "filename": "old-outline.txt",
            "chapter_name": "旧考纲",
            "file_path": str(source_path),
            "file_type": "txt",
        }
        self.knowledge_base._start_workbench_upload_background = started_jobs.append
        try:
            job = self.knowledge_base._queue_workbench_reprocess(1, source)
        finally:
            self.knowledge_base._start_workbench_upload_background = original_start

        self.assertEqual(job["material_id"], material["id"])
        self.assertTrue(job["replace_existing"])
        self.assertEqual(len(started_jobs), 1)
        with sqlite3.connect(self.db_path) as conn:
            status = conn.execute(
                "SELECT processing_status FROM user_materials WHERE id=?",
                (material["id"],),
            ).fetchone()[0]
            remaining_knowledge = conn.execute(
                "SELECT COUNT(*) FROM user_knowledge WHERE material_id=?",
                (material["id"],),
            ).fetchone()[0]
        self.assertEqual(status, "pending")
        self.assertEqual(remaining_knowledge, 1)

    def test_wrongbook_html_never_receives_service_api_key(self):
        app_source = (PROJECT_DIR / "app.py").read_text(encoding="utf-8")

        self.assertNotIn("WB_API_KEY", app_source)
        self.assertNotIn("_start_wb_save_server", app_source)

    def test_page_exposes_dedicated_syllabus_reader(self):
        app = self._run_app()

        self.assertIn("docx", self.knowledge_base.SUPPORTED_SYLLABUS_FILE_TYPES)
        self.assertIn("docx", self.knowledge_base.SUPPORTED_MATERIAL_FILE_TYPES)
        source_mode = _find_by_label(app.radio, "添加方式")
        self.assertEqual(source_mode.value, "读取大纲")
        self.assertTrue(any(item.label == "上传考试大纲" for item in app.get("file_uploader")))
        self.assertTrue(any(item.label == "背诵条目数量" for item in app.slider))
        self.assertTrue(
            any(item.label == "读取大纲并生成背诵内容" for item in app.button)
        )

        source_mode.set_value("添加资料").run()
        if app.exception:
            raise AssertionError(app.exception)
        self.assertTrue(any(item.label == "上传资料" for item in app.get("file_uploader")))
        self.assertTrue(any(item.label == "添加来源" for item in app.button))
        self.assertFalse(any(item.label == "本地资料文件夹" for item in app.text_input))
        self.assertFalse(any(item.label == "本机资料文件夹" for item in app.text_input))

    def test_syllabus_upload_uses_dedicated_background_worker(self):
        started_jobs = []
        original_start = self.knowledge_base._start_syllabus_memorization_background

        class FakeUpload:
            name = "历史考试大纲.txt"

            def getvalue(self):
                return "中国古代史：商鞅变法；中国近现代史：洋务运动。".encode("utf-8")

        self.knowledge_base._start_syllabus_memorization_background = started_jobs.append
        try:
            job = self.knowledge_base._queue_syllabus_upload(
                1,
                "历史学基础",
                FakeUpload(),
                max_points=80,
            )
        finally:
            self.knowledge_base._start_syllabus_memorization_background = original_start
            Path(job["file_path"]).unlink(missing_ok=True)

        self.assertEqual(len(started_jobs), 1)
        self.assertEqual(job["max_points"], 80)
        self.assertTrue(job["chapter_name"].startswith("学校考纲 - "))

    def test_syllabus_worker_saves_multiple_expanded_points(self):
        from repositories.material_repo import create_material

        source_path = self.temp_dir / "history-outline.txt"
        source_path.write_text(
            "中国古代史：商鞅变法；中国近现代史：洋务运动。",
            encoding="utf-8",
        )
        with sqlite3.connect(self.db_path) as conn:
            material = create_material(
                conn,
                user_id=1,
                subject="历史学基础",
                filename=source_path.name,
                chapter_name="学校考纲 - 历史",
                file_path=str(source_path),
                file_type="txt",
            )
            conn.commit()
        task = self.task_service.create_task(
            user_id=1,
            subject="历史学基础",
            chapter_name="学校考纲 - 历史",
            filename=source_path.name,
            material_id=material["id"],
        )
        job = {
            "user_id": 1,
            "material_id": material["id"],
            "task_id": task.task_id,
            "subject": "历史学基础",
            "chapter_name": "学校考纲 - 历史",
            "filename": source_path.name,
            "file_path": str(source_path),
            "file_type": "txt",
            "max_points": 60,
        }
        fake_points = [
            {
                "knowledge_name": "商鞅变法",
                "knowledge_type": "背诵知识点",
                "subject": "历史学基础",
                "chapter_name": "中国古代史",
                "core_definition": "商鞅变法围绕政治、经济和军事制度展开，推动秦国建立较系统的中央集权制度并增强国力。",
                "keywords": ["秦国", "变法"],
                "is_ai_expansion": True,
            },
            {
                "knowledge_name": "洋务运动",
                "knowledge_type": "背诵知识点",
                "subject": "历史学基础",
                "chapter_name": "中国近现代史",
                "core_definition": "洋务运动以自强、求富为目标兴办近代军事和民用企业，并推动近代教育、海军和技术传播。",
                "keywords": ["自强", "求富"],
                "is_ai_expansion": True,
            },
        ]
        original_generate = self.knowledge_base.generate_syllabus_memorization_points
        self.knowledge_base.generate_syllabus_memorization_points = (
            lambda *args, **kwargs: (fake_points, [])
        )
        os.environ["AI_API_KEY"] = "test-key"
        try:
            self.knowledge_base._process_syllabus_memorization_job(job)
        finally:
            self.knowledge_base.generate_syllabus_memorization_points = original_generate

        with sqlite3.connect(self.db_path) as conn:
            source_state = conn.execute(
                """SELECT processing_status, process_method, knowledge_count, error_message
                   FROM user_materials WHERE id=?""",
                (material["id"],),
            ).fetchone()
            self.assertEqual(source_state, ("done", "syllabus_memorization_ai", 2, ""))
            rows = conn.execute(
                """SELECT knowledge_name, chapter_name FROM user_knowledge
                   WHERE material_id=? ORDER BY knowledge_name""",
                (material["id"],),
            ).fetchall()
        self.assertEqual(
            rows,
            [("商鞅变法", "中国古代史"), ("洋务运动", "中国近现代史")],
        )
        saved_task = self.task_service.load_task(task.task_id)
        self.assertEqual(saved_task.status, "done")
        self.assertIn("共 2 个条目", saved_task.notes[-1])

if __name__ == "__main__":
    unittest.main()
