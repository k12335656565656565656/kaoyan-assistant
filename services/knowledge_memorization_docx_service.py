from __future__ import annotations

import io
import json
import re
from collections import OrderedDict
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


_GROUP_COLORS = ("2563A6", "B65C22", "2E7D63", "8054A0", "9A6700", "0F766E")
_NOISE_MARKERS = (
    "扫码",
    "加微信",
    "qq群",
    "更多计算机考研资料",
    "候选知识点",
    "未命名知识点",
)


def _safe_text(value: Any, fallback: str = "") -> str:
    text = str(value or "").strip()
    return text or fallback


def _is_noisy(value: str) -> bool:
    compact = "".join(_safe_text(value).lower().split())
    return any(marker in compact for marker in _NOISE_MARKERS)


def _json_list(value: Any) -> list[str]:
    if isinstance(value, (list, tuple, set)):
        items = list(value)
    elif not value:
        items = []
    else:
        try:
            parsed = json.loads(value)
        except (TypeError, ValueError, json.JSONDecodeError):
            parsed = None
        if isinstance(parsed, list):
            items = parsed
        else:
            items = re.split(r"[，,；;、\n]+", str(value))
    cleaned = [_safe_text(item).rstrip("。") for item in items if _safe_text(item)]
    return [item for item in cleaned if not _is_noisy(item)]


def _set_run_font(run, *, size: float = 10, bold: bool | None = None, color: str | None = None):
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_style_font(style, *, size: float, bold: bool = False, color: str = "273142"):
    style.font.name = "Microsoft YaHei"
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_border(cell, color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "5")
        node.set(qn("w:color"), color)
        borders.append(node)


def _set_cell_margins(cell, *, top=100, start=150, bottom=95, end=150):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)


def _add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def _extract_core_sections(value: Any) -> dict[str, str]:
    text = _safe_text(value)
    sections: dict[str, str] = {}
    for block in re.split(r"\n\s*\n", text):
        clean = block.strip()
        if not clean:
            continue
        if "：" in clean:
            label, content = clean.split("：", 1)
            if label.strip() in {"核心理解", "常见考法", "易错提醒", "掌握标准", "关联知识"}:
                sections[label.strip()] = content.strip()
                continue
        sections.setdefault("核心理解", clean)
    return sections


def _group_points(points: list[dict]) -> OrderedDict[str, list[dict]]:
    groups: OrderedDict[str, list[dict]] = OrderedDict()
    for point in points:
        title = _safe_text(point.get("knowledge_name"))
        if not title or _is_noisy(title):
            continue
        group_name = _safe_text(point.get("chapter_name"), "未分类")
        groups.setdefault(group_name, []).append(point)
    return groups


def _field_items(point: dict, plain_key: str, json_key: str) -> list[str]:
    return _json_list(point.get(plain_key) or point.get(json_key))


def _add_labeled_paragraph(cell, label: str, value: str, color: str):
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    label_run = paragraph.add_run(f"{label}：")
    _set_run_font(label_run, size=9.5, bold=True, color=color)
    value_run = paragraph.add_run(_safe_text(value))
    _set_run_font(value_run, size=9.5)


def _add_entry(document: Document, point: dict, number: str, color: str):
    title = _safe_text(point.get("knowledge_name"), "知识点")
    heading = document.add_paragraph(style="Heading 2")
    heading_run = heading.add_run(f"{number}  {title}")
    _set_run_font(heading_run, size=12.5, bold=True, color=color)

    card = document.add_table(rows=1, cols=1)
    card.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = card.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _set_cell_border(cell, "D8E0EA")
    _set_cell_margins(cell)
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    core_sections = _extract_core_sections(point.get("core_definition") or point.get("content"))
    core = core_sections.get("核心理解") or ""
    if core:
        _add_labeled_paragraph(cell, "核心内容", core, color)

    keywords = _field_items(point, "keywords", "keywords_json")
    if keywords:
        _add_labeled_paragraph(cell, "关键词", "；".join(keywords) + "。", color)

    related = _field_items(point, "related_concepts", "related_concepts_json")
    related_text = "；".join(related) if related else core_sections.get("关联知识", "")
    if related_text:
        _add_labeled_paragraph(cell, "关联知识", related_text.rstrip("。") + "。", color)

    exam = _field_items(point, "exam_question_styles", "exam_question_styles_json")
    exam_text = "；".join(exam) if exam else core_sections.get("常见考法", "")
    if exam_text:
        _add_labeled_paragraph(cell, "常见考法", exam_text.rstrip("。") + "。", color)

    pitfalls = _field_items(point, "pitfalls", "pitfalls_json")
    pitfall_text = "；".join(pitfalls) if pitfalls else core_sections.get("易错提醒", "")
    if pitfall_text:
        _add_labeled_paragraph(cell, "易错点", pitfall_text.rstrip("。") + "。", color)

    example = _safe_text(point.get("example_or_application"))
    if example:
        _add_labeled_paragraph(cell, "例题与应用", example, color)

    mastery = core_sections.get("掌握标准", "")
    if mastery:
        _add_labeled_paragraph(cell, "掌握标准", mastery, color)

    _add_labeled_paragraph(
        cell,
        "背诵自测",
        f"不看正文，完整讲清“{title}”，再说出一种常见考法和一个易错点。",
        color,
    )


def build_knowledge_memorization_docx(points: list[dict], *, subject: str) -> bytes:
    """Build a clean A4 DOCX intended for students to print and memorize."""
    groups = _group_points(list(points or []))
    if not groups:
        raise ValueError("没有可导出的知识条目。")

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)
    section.different_first_page_header_footer = True

    styles = document.styles
    _set_style_font(styles["Normal"], size=10)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    styles["Normal"].paragraph_format.line_spacing = 1.2
    _set_style_font(styles["Title"], size=26, bold=True, color="172033")
    _set_style_font(styles["Subtitle"], size=12, color="667085")
    _set_style_font(styles["Heading 1"], size=18, bold=True, color="27366F")
    styles["Heading 1"].paragraph_format.keep_with_next = True
    _set_style_font(styles["Heading 2"], size=12.5, bold=True, color="172033")
    styles["Heading 2"].paragraph_format.keep_with_next = True

    document.core_properties.title = f"{_safe_text(subject, '专业课')}背诵手册"
    document.core_properties.subject = "专业课背诵打印版"
    document.core_properties.author = "考研学习助手"

    settings = document.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    cover_spacer = document.add_paragraph()
    cover_spacer.paragraph_format.space_after = Pt(72)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"{_safe_text(subject, '专业课')}背诵手册")
    _set_run_font(title_run, size=26, bold=True)
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("专业课知识点打印版")
    _set_run_font(subtitle_run, size=12, color="667085")
    subtitle.add_run().add_break(WD_BREAK.PAGE)

    document.add_paragraph("目录", style="Heading 1")
    for group_index, (group_name, group_points) in enumerate(groups.items(), start=1):
        group_line = document.add_paragraph()
        group_run = group_line.add_run(f"{group_index}. {group_name}")
        _set_run_font(group_run, size=11, bold=True, color=_GROUP_COLORS[(group_index - 1) % len(_GROUP_COLORS)])
        for item_index, point in enumerate(group_points, start=1):
            item_line = document.add_paragraph()
            item_line.paragraph_format.left_indent = Cm(0.6)
            item_line.paragraph_format.space_after = Pt(1)
            item_run = item_line.add_run(f"{group_index}.{item_index}  {_safe_text(point.get('knowledge_name'))}")
            _set_run_font(item_run, size=9)
    for group_index, (group_name, group_points) in enumerate(groups.items(), start=1):
        color = _GROUP_COLORS[(group_index - 1) % len(_GROUP_COLORS)]
        heading = document.add_paragraph(style="Heading 1")
        heading_run = heading.add_run(f"第{group_index}部分  {group_name}")
        _set_run_font(heading_run, size=18, bold=True, color=color)
        for item_index, point in enumerate(group_points, start=1):
            _add_entry(document, point, f"{group_index}.{item_index}", color)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run(f"{_safe_text(subject, '专业课')}背诵手册")
    _set_run_font(header_run, size=8, color="7A869A")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_prefix = footer.add_run("第 ")
    _set_run_font(page_prefix, size=8, color="667085")
    _add_page_field(footer)
    page_suffix = footer.add_run(" 页")
    _set_run_font(page_suffix, size=8, color="667085")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
